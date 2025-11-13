#!/usr/bin/env python3
"""
Convert Executive Business Overview markdown to professional Word document.
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_professional_document():
    """Create a professionally formatted Word document from markdown."""

    # Read the markdown file
    md_path = Path(__file__).parent.parent / 'docs' / 'EXECUTIVE_BUSINESS_OVERVIEW.md'
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create document
    doc = Document()

    # Set up styles
    setup_styles(doc)

    # Parse and add content
    lines = content.split('\n')
    i = 0
    current_table_data = []
    in_table = False
    in_blockquote = False

    while i < len(lines):
        line = lines[i].rstrip()

        # Skip table of contents (we'll regenerate it)
        if '## Table of Contents' in line:
            i += 1
            while i < len(lines) and not lines[i].startswith('##'):
                i += 1
            continue

        # Handle horizontal rules
        if line.strip() == '---':
            doc.add_paragraph('_' * 50)
            i += 1
            continue

        # Handle blockquotes
        if line.startswith('> '):
            text = line[2:]
            p = doc.add_paragraph(text, style='Quote')
            i += 1
            continue

        # Handle headings
        if line.startswith('#'):
            level = len(re.match(r'^#+', line).group())
            text = line.lstrip('#').strip()

            if level == 1:
                p = doc.add_heading(text, level=0)
                p.runs[0].font.color.rgb = RGBColor(10, 37, 64)  # Deep Navy
            elif level == 2:
                p = doc.add_heading(text, level=1)
                p.runs[0].font.color.rgb = RGBColor(10, 37, 64)
            elif level == 3:
                p = doc.add_heading(text, level=2)
                p.runs[0].font.color.rgb = RGBColor(45, 106, 79)  # Forest Green
            else:
                p = doc.add_heading(text, level=3)

            i += 1
            continue

        # Handle tables
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                current_table_data = []

            # Parse table row
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            current_table_data.append(cells)

            i += 1

            # Check if next line is separator or not a table
            if i >= len(lines) or '|' not in lines[i]:
                # Create table
                if len(current_table_data) > 1:
                    # Filter out separator rows
                    table_data = [row for row in current_table_data if not all(re.match(r'^:?-+:?$', cell) for cell in row)]

                    if table_data:
                        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                        table.style = 'Light Grid Accent 1'

                        for row_idx, row_data in enumerate(table_data):
                            for col_idx, cell_data in enumerate(row_data):
                                cell = table.rows[row_idx].cells[col_idx]
                                cell.text = cell_data

                                # Bold header row
                                if row_idx == 0:
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.bold = True
                                            run.font.color.rgb = RGBColor(255, 255, 255)
                                    cell._element.get_or_add_tcPr().append(
                                        parse_xml(r'<w:shd {} w:fill="0A2540"/>'.format(nsdecls('w')))
                                    )

                        doc.add_paragraph()  # Add space after table

                in_table = False
                current_table_data = []
            continue

        # Handle bullet lists
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            # Parse bold text
            text = parse_inline_formatting(text)
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
            i += 1
            continue

        # Handle numbered lists
        if re.match(r'^\d+\. ', line.strip()):
            text = re.sub(r'^\d+\. ', '', line.strip())
            text = parse_inline_formatting(text)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)
            i += 1
            continue

        # Handle regular paragraphs
        if line.strip():
            text = parse_inline_formatting(line)
            p = doc.add_paragraph()
            add_formatted_text(p, text)
        else:
            # Empty line - add small space
            doc.add_paragraph()

        i += 1

    # Save document
    output_path = Path(__file__).parent.parent / 'docs' / 'RostraCore_Executive_Business_Overview.docx'
    doc.save(output_path)
    print(f"✓ Professional Word document created: {output_path}")
    print(f"  File size: {output_path.stat().st_size / 1024:.1f} KB")
    return output_path

def setup_styles(doc):
    """Set up professional document styles."""
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Quote style
    if 'Quote' not in [s.name for s in doc.styles]:
        quote_style = doc.styles.add_style('Quote', WD_STYLE_TYPE.PARAGRAPH)
        quote_style.font.italic = True
        quote_style.font.color.rgb = RGBColor(45, 106, 79)
        quote_style.font.size = Pt(12)

def parse_inline_formatting(text):
    """Parse inline markdown formatting."""
    # Don't process if empty
    if not text:
        return text

    # Keep markdown markers for processing
    return text

def add_formatted_text(paragraph, text):
    """Add text with inline formatting to paragraph."""
    if not text:
        return

    # Split by bold markers
    parts = re.split(r'(\*\*.*?\*\*)', text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            # Italic text
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
        else:
            # Regular text - handle inline code
            code_parts = re.split(r'(`[^`]+`)', part)
            for code_part in code_parts:
                if code_part.startswith('`') and code_part.endswith('`'):
                    run = paragraph.add_run(code_part[1:-1])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                else:
                    paragraph.add_run(code_part)

def parse_xml(xml_string):
    """Helper function for table styling."""
    from docx.oxml import parse_xml as px
    return px(xml_string)

def nsdecls(*prefixes):
    """Helper function for XML namespaces."""
    from docx.oxml.ns import nsdecls as ns
    return ns(*prefixes)

if __name__ == '__main__':
    create_professional_document()
